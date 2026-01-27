import pandas as pd
from docx import Document
from docx.shared import RGBColor
import win32com.client as win32
import shutil
import os
import datetime
import locale

def lire_trois_premieres_colonnes_excel(chemin_fichier_excel, nom_feuille):
    liste_artisans = []
    # Lire le fichier Excel avec pandas, en sautant les 14 premières lignes
    df = pd.read_excel(chemin_fichier_excel, sheet_name=nom_feuille, skiprows=13)

    # Extraire les trois premières colonnes
    trois_premieres_colonnes = df.iloc[:, :3]  # Sélectionne les trois premières colonnes
    trois_premieres_colonnes = trois_premieres_colonnes.values.tolist()

    # changer le nom de la feuille si c'est CCTP pour pointer sur la feuille AVP correspondante
    # car seules les feuilles AVP contiennent la liste des artisans
    if "CCTP" in nom_feuille:   
        nom_feuille = nom_feuille.replace("CCTP", "AVP")
        df = pd.read_excel(chemin_fichier_excel, sheet_name=nom_feuille, skiprows=13)
        troisieme_colonne = df.iloc[:, 2]  # Sélectionne la troisième colonne
        troisieme_colonne = troisieme_colonne.dropna().unique()  # Obtenir les valeurs uniques non nulles
        liste_artisans = troisieme_colonne.tolist()

    # Retourner les données sous forme de liste de listes
    return trois_premieres_colonnes, liste_artisans

def ajouter_dans_fichier_word(word_en_sortie, donnees):
   
    # Charger le document Word existant
    doc = Document(word_en_sortie)

    # Ajouter chaque élément de la première colonne au document Word
    for valeur in donnees:
        v0 = str(valeur[0])
        if 'Titre 1' in v0:
            doc.add_paragraph(str(valeur[1]), style='Heading 1')
        elif 'Titre 2' in v0:
            doc.add_paragraph(str(valeur[1]), style='Heading 2')
        elif 'Titre 3' in v0:
            doc.add_paragraph(str(valeur[1]), style='Heading 3')
        elif 'Normal' in v0:
            doc.add_paragraph(str(valeur[1]), style='Normal')
        elif 'jaune' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.color.rgb = RGBColor(255, 255, 0)  # Changer la couleur du texte en jaune
            if 'barre' in v0 or 'barré' in v0:
                run.font.strike = True  # Appliquer le style barré
        elif 'bleu' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.color.rgb = RGBColor(0, 0, 255)  # Changer la couleur du texte en bleu
            if 'barre' in v0 or 'barré' in v0:
                run.font.strike = True  # Appliquer le style barré
        elif 'vert' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.color.rgb = RGBColor(20, 148, 20)  # Changer la couleur du texte en vert
            if 'barre' in v0 or 'barré' in v0:
                run.font.strike = True  # Appliquer le style barré
        elif 'violet' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.color.rgb = RGBColor(128, 28, 248)  # Changer la couleur du texte en violet
            if 'barre' in v0 or 'barré' in v0:
                run.font.strike = True  # Appliquer le style barré
        elif 'orange' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.color.rgb = RGBColor(233, 109, 0)  # Changer la couleur du texte en orange
            if 'barre' in v0 or 'barré' in v0:
                run.font.strike = True  # Appliquer le style barré
        elif 'barre' in v0 or 'barré' in v0:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[1]))
            run.font.strike = True  # Appliquer le style barré
        else:
            if str(valeur[1]) != 'nan':  # Vérifier si la deuxième colonne n'est pas vide
                doc.add_paragraph(str(valeur[1]), style='Normal')
        if (str(valeur[2]) != 'nan') : # Vérifier si la troisième colonne n'est pas vide
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(valeur[2]))
            run.bold = True  # Mettre le texte en gras
            run.font.color.rgb = RGBColor(255, 0, 0)  # Changer la couleur du texte en rouge

    # Sauvegarder le document Word
    doc.save(word_en_sortie)



def ajouter_liste_artisans(word_en_sortie, liste_artisans):
    # Charger le document Word existant
    doc = Document(word_en_sortie)
 
    doc.add_page_break()
    doc.add_paragraph("Signatures et tampons", style='Heading 1')
    doc.add_paragraph("Fait à Balma                                     le : ", style='Normal')
    table = doc.add_table(rows=len(liste_artisans)+2, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'Représenté par'
    hdr_cells[2].text = 'Signature'
    hdr_cells[3].text = 'Cachet'

    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Le maitre d\'ouvrage\n\n'
    # Remplir les noms des entreprises dans la première colonne
    for artisan_index, artisan in enumerate(liste_artisans):
        hdr_cells = table.rows[artisan_index+2].cells
        hdr_cells[0].text = f"{artisan}\n\n"

    # Sauvegarder le document Word
    doc.save(word_en_sortie)


def lire_donnees_client_excel(param_chemin_fichier_excel, param_nom_feuille):
    # Lire le fichier Excel avec pandas
    df = pd.read_excel(param_chemin_fichier_excel, sheet_name=param_nom_feuille, usecols="D:E")

     # Créer une liste pour stocker les valeurs des cellules D10 à D12
    valeurs = []

    #Lire les valeurs de la deuxième colonne pour les lignes 4 à 8 (E10 à E14 dans Excel)
    for index in range(1, 5):  
        valeur = df.iloc[index, 1]  # Colonne E (index 1)
        valeurs.append(valeur)

    # Recupérer le numero de version dans le nom de la feuille
    param_sheet_split = param_nom_feuille.split(" ")
    if len(param_sheet_split) > 1:
        valeurs.append(param_sheet_split[0])  # Ajouter le type de document (AVP ou CCTP)
        valeurs.append(param_sheet_split[1])  # Ajouter la version
    else:
        valeurs.append(param_sheet_split[0])  # Ajouter le type de document (AVP ou CCTP)
        valeurs.append("1")  # Valeur par défaut si aucune version n'est trouvée

    print(f"Les valeurs extraites du fichier Excel sont : {valeurs}")
    return valeurs

def supprimer_dossier_temp(nom_dossier):
    # Obtenir le chemin du répertoire %TEMP%
    temp_dir = os.environ.get('TEMP')

    # Construire le chemin complet du dossier à supprimer
    chemin_dossier = os.path.join(temp_dir, nom_dossier)

    # Vérifier si le dossier existe
    if os.path.exists(chemin_dossier):
        try:
            # Supprimer le dossier et tout son contenu
            shutil.rmtree(chemin_dossier)
            print(f"Le dossier {chemin_dossier} a été supprimé avec succès.")
        except Exception as e:
            print(f"Une erreur est survenue lors de la suppression du dossier : {e}")
    else:
        print(f"Le dossier {chemin_dossier} n'existe pas.")

def mise_a_jour_signets(word_en_sortie, donnees_client):
    
    try:
        # Créer une instance de Word
        word_app = win32.gencache.EnsureDispatch('Word.Application')
        word_app.Visible = False  # Ne pas rendre Word visible (pour un traitement en arrière-plan)
    except Exception as e:
        # supprimer le cache et réessayer
        supprimer_dossier_temp('gen_py')
        word_app = win32.gencache.EnsureDispatch('Word.Application')   
        word_app.Visible = False  # Ne pas rendre Word visible (pour un traitement en arrière-plan)
    
    try:
        # Ouvrir le document Word
        doc = word_app.Documents.Open(word_en_sortie)

        # Vérifier si le signet "AdresseProjet" existe et le remplir
        if doc.Bookmarks.Exists("AdresseProjet"):
            doc.Bookmarks("AdresseProjet").Range.Text = str(donnees_client[2])  
        if doc.Bookmarks.Exists("MaitreOuvrage"):
            doc.Bookmarks("MaitreOuvrage").Range.Text = str(donnees_client[1])  
        if doc.Bookmarks.Exists("VersionDocument"):
            doc.Bookmarks("VersionDocument").Range.Text = str(donnees_client[5])  # Version du document
        if doc.Bookmarks.Exists("CoordonneesClient"):
            doc.Bookmarks("CoordonneesClient").Range.Text = str(donnees_client[3])  # Coordonnées du client
        
        # On modifie le titre du document en fonction du type de document
        if doc.Bookmarks.Exists("TypeDocument"):
            if str(donnees_client[4]) == "AVP":
                doc.Bookmarks("TypeDocument").Range.Text = "Etude d'Avant-Projet \n(AVP)"
            elif str(donnees_client[4]) == "CCTP":
                doc.Bookmarks("TypeDocument").Range.Text = "Cahier des Clauses Techniques Particulières \n(CCTP)"
        
        # On ajout la date du jour au format jj mois aaaa
        if doc.Bookmarks.Exists("DateGen"):
            # Définir la locale en français pour obtenir le nom du mois en français
            locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
            date_aujourdhui = datetime.datetime.now()
            date_formatee = date_aujourdhui.strftime("%d %B %Y")
            doc.Bookmarks("DateGen").Range.Text = date_formatee
        
        # Mettre à jour toutes les tables des matières
        for table in doc.TablesOfContents:
            table.Update()

        # Sauvegarder et fermer le document Word
        doc.Save()
        doc.Close()
        #print(f"Zeus")
    except Exception as e:
        print(f"Une erreur est survenue : {e}")
    finally:
        # Quitter Word
        word_app.Quit()

if __name__ == "__main__":
    if len(os.sys.argv) != 3:
        print("Usage: python xlstodocx.py <chemin_fichier_excel> <nom_feuille>")
        os.sys.exit(1)

    param_chemin_fichier_excel = os.sys.argv[1]
    param_nom_feuille = os.sys.argv[2]
    
    # Chemins des fichiers
    word_en_entree = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'DSTest.docx')  # Remplacez par le chemin de votre fichier Word

    donnees_client = lire_donnees_client_excel(param_chemin_fichier_excel, param_nom_feuille)
    print(f"Les données client sont : {donnees_client}")
    try:
        # Vérifier si le fichier source existe
        if not os.path.exists(word_en_entree):
            raise FileNotFoundError(f"Le fichier source {word_en_entree} n'existe pas.")

        # Lire les informations client du fichier Excel
        donnees_client = lire_donnees_client_excel(param_chemin_fichier_excel, param_nom_feuille)
        # Définir le chemin du fichier Word de sortie qui va être dans le même dossier que le fichier Excel
        word_en_sortie = os.path.dirname(param_chemin_fichier_excel) + "\\" + str(donnees_client[1]) + "_" + str(donnees_client[4]) + "_V" + str(donnees_client[5]) + ".docx"
        # word_en_sortie = os.getcwd() + "\\" + str(donnees_client[0]) + "_" + str(donnees_client[4]) + "_V" + str(donnees_client[5]) + ".docx"
        # Copier le fichier
        shutil.copy(word_en_entree, word_en_sortie)
        print(f"Le fichier a été copié de {word_en_entree} vers {word_en_sortie}")

    except Exception as e:
        print(f"Une erreur est survenue : {e}")   



    # Lire les 3 premières colonnes du fichier Excel et obtenir la liste des artisans
    donnees_premiere_colonne, liste_artisans = lire_trois_premieres_colonnes_excel(param_chemin_fichier_excel, param_nom_feuille)

    # Creer la liste des artisans uniques
    print(liste_artisans)

    # Afficher les valeurs
    #print(f"Les valeurs des cellules D10 à D12 sont : {donnees_client}")


    # Ajouter les données dans le fichier Word existant
    ajouter_dans_fichier_word(word_en_sortie, donnees_premiere_colonne)

    # Ajouter la liste des artisans si le document est un CCTP
    if "CCTP" in word_en_sortie:
        ajouter_liste_artisans(word_en_sortie, liste_artisans)
        


    # Mettre à jour les signets du document Word
    mise_a_jour_signets(word_en_sortie, donnees_client)

    
    print(f"Les données ont été ajoutées à {word_en_sortie}")
